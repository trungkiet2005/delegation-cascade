"""Build the separate submission files Applied Mathematics and Computation asks for.

The journal wants separate items uploaded alongside the manuscript:

* highlights, as a separate editable file whose name contains "highlights";
* the declaration of competing interests, as a .doc/.docx file.

    python paper-amc/scripts/make_submission_files.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

OUT = Path(__file__).resolve().parents[1]

TITLE = (
    "Delegation cascades in evolutionary games: stable computation with "
    "strategic depth"
)

HIGHLIGHTS = [
    "Strategic depth creates a liability shelter in evolutionary games",
    "Maximum-shifted O(Z) fixation sums stabilise the reduced Markov chain",
    "Joint erosion and attribution loss raise unsafe frequency from 0 to 0.322",
    "A floor closely traces the depth-cap frontier over its attainable range",
]

AUTHORS = (
    "Trung-Kiet Huynh, Dao-Sy Duy-Minh, Chi-Nguyen Tran, "
    "Nguyen Lam Phu Quy, Pham Phu Hoa"
)

COMPETING_INTERESTS = (
    "The authors declare that they have no known competing financial interests "
    "or personal relationships that could have appeared to influence the work "
    "reported in this paper."
)

def _document() -> Document:
    doc = Document()
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        # python-docx's default template sets bestFit but omits the percentage
        # required by the ECMA-376 schema used in the submission validator.
        zoom.set(qn("w:percent"), "100")
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    return doc


def write_highlights() -> None:
    doc = _document()
    doc.add_heading("Highlights", level=1)
    doc.add_paragraph(TITLE).runs[0].bold = True
    author_paragraph = doc.add_paragraph()
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    author_paragraph.add_run(AUTHORS).font.size = Pt(9.5)
    for item in HIGHLIGHTS:
        assert len(item) <= 85, "highlight over 85 characters: %r" % item
        doc.add_paragraph(item, style="List Bullet")
    doc.save(OUT / "highlights.docx")

    (OUT / "highlights.txt").write_text(
        "Highlights\n\n%s\n\n" % TITLE
        + "\n".join("- " + h for h in HIGHLIGHTS)
        + "\n",
        encoding="utf-8",
    )


def write_declarations() -> None:
    doc = _document()
    doc.add_heading("Declaration of competing interest", level=1)
    doc.add_paragraph(TITLE).runs[0].italic = True
    doc.add_paragraph(COMPETING_INTERESTS)
    doc.save(OUT / "declaration_of_interest.docx")


def main() -> None:
    write_highlights()
    write_declarations()
    for name in ("highlights.docx", "highlights.txt", "declaration_of_interest.docx"):
        print(name, (OUT / name).stat().st_size, "bytes")


if __name__ == "__main__":
    main()
